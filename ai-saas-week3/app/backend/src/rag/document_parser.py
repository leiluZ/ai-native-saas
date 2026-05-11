import logging
import asyncio
import json
import mimetypes
from typing import List
from datetime import datetime
from pathlib import Path


from .schemas import ParsedDocument, ParseError

logger = logging.getLogger(__name__)


class DocumentParser:
    MIME_PARSERS = {
        "application/pdf": "_parse_pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_parse_docx",
        "text/html": "_parse_html",
        "text/markdown": "_parse_markdown",
        "text/plain": "_parse_text",
    }

    def __init__(self):
        self.error_log: List[ParseError] = []

    async def parse_files(self, file_paths: List[str]) -> dict:
        logger.info(
            f"[DocumentParser] Starting batch parse - files_count={len(file_paths)}"
        )
        tasks = [self.parse_file(f) for f in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        success = []
        errors = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                errors.append(
                    ParseError(
                        file_path=file_paths[i],
                        error=str(result),
                        timestamp=datetime.now().isoformat(),
                    )
                )
                logger.warning(
                    f"[DocumentParser] Batch parse failed - file='{file_paths[i]}', error={str(result)}"
                )
            else:
                success.append(result.model_dump())

        logger.info(
            f"[DocumentParser] Batch parse complete - success={len(success)}, failed={len(errors)}"
        )

        return {"success": success, "errors": errors}

    async def parse_file(self, file_path: str) -> ParsedDocument:
        logger.info(f"[DocumentParser] Parsing file - path='{file_path}'")

        ext = Path(file_path).suffix.lower()
        ext_to_mime = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".html": "text/html",
            ".md": "text/markdown",
            ".txt": "text/plain",
        }

        mime_type = ext_to_mime.get(ext)
        if mime_type is None:
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type is None:
                mime_type = "text/plain"
                logger.debug(
                    "[DocumentParser] Unknown MIME type, defaulting to text/plain"
                )

        parser_method = self.MIME_PARSERS.get(mime_type, "_parse_text")
        logger.debug(
            f"[DocumentParser] Using parser - method='{parser_method}', mime_type='{mime_type}'"
        )

        try:
            content = await getattr(self, parser_method)(file_path)
            cleaned = self._clean_content(content)
            metadata = {
                "source": file_path,
                "mime_type": mime_type,
                "parsed_at": datetime.now().isoformat(),
            }
            logger.info(
                f"[DocumentParser] Parse success - file='{file_path}', raw_length={len(content)}, "
                f"cleaned_length={len(cleaned)}, mime_type='{mime_type}'"
            )
            return ParsedDocument(content=cleaned, metadata=metadata)
        except Exception as e:
            logger.error(
                f"[DocumentParser] Parse failed - file='{file_path}', error={str(e)}"
            )
            error = ParseError(
                file_path=file_path, error=str(e), timestamp=datetime.now().isoformat()
            )
            self.error_log.append(error)
            self._save_error_log()
            raise

    async def parse_directory(
        self, directory: str, recursive: bool = True
    ) -> List[ParsedDocument]:
        path = Path(directory)
        pattern = "**/*" if recursive else "*"
        files = [str(f) for f in path.glob(pattern) if f.is_file()]

        logger.info(
            f"[DocumentParser] Parsing directory - path='{directory}', recursive={recursive}, files_count={len(files)}"
        )

        tasks = [self.parse_file(f) for f in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        parsed = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                error = ParseError(
                    file_path=files[i],
                    error=str(result),
                    timestamp=datetime.now().isoformat(),
                )
                self.error_log.append(error)
                logger.warning(
                    f"[DocumentParser] Directory parse failed - file='{files[i]}', error={str(result)}"
                )
            else:
                parsed.append(result)

        if self.error_log:
            self._save_error_log()

        logger.info(
            f"[DocumentParser] Directory parse complete - parsed={len(parsed)}, errors={len(self.error_log)}"
        )

        return parsed

    async def _parse_pdf(self, file_path: str) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            return await self._parse_with_unstructured(file_path)

    async def _parse_docx(self, file_path: str) -> str:
        try:
            from docx import Document

            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            return await self._parse_with_unstructured(file_path)

    async def _parse_html(self, file_path: str) -> str:
        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        return soup.get_text(separator=" ", strip=True)

    async def _parse_markdown(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    async def _parse_text(self, file_path: str) -> str:
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise RuntimeError(f"无法解码文件，尝试了编码: {encodings}")

    async def _parse_with_unstructured(self, file_path: str) -> str:
        try:
            from unstructured.partition.auto import partition

            elements = partition(filename=file_path)
            return "\n".join([str(el) for el in elements])
        except ImportError:
            # 如果 unstructured 不可用，检查是否为二进制文件
            ext = Path(file_path).suffix.lower()
            if ext in [".pdf", ".docx"]:
                raise RuntimeError(
                    f"无法解析 {ext} 文件：缺少 pdfplumber/unstructured 依赖"
                )
            return await self._parse_text(file_path)

    def _clean_content(self, content: str) -> str:
        lines = content.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if any(marker in line for marker in ["---PAGE---", "[PAGE]", "Page \\d+"]):
                continue
            if line.isdigit() and len(line) < 5:
                continue
            cleaned_lines.append(line)

        text = "\n".join(cleaned_lines)
        text = "\n".join(line for line in text.split("\n") if line.strip())

        return text

    def _save_error_log(self):
        try:
            import tempfile

            log_path = Path(tempfile.gettempdir()) / "rag_error_log.json"
            with open(log_path, "w", encoding="utf-8") as f:
                json.dump(
                    [e.model_dump() for e in self.error_log],
                    f,
                    ensure_ascii=False,
                    indent=2,
                )
            logger.debug(f"[DocumentParser] Error log saved to '{log_path}'")
        except Exception as e:
            logger.warning(f"[DocumentParser] Failed to save error log: {str(e)}")
